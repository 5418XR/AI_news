import os
import json
import requests
import tempfile
import docx2txt
from flask import Flask, render_template, request, jsonify, session, send_file
from docx import Document
from io import BytesIO

app = Flask(__name__)
app.secret_key = os.urandom(24)  # 用于会话加密

# DeepSeek API配置
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"

@app.route('/')
def index():
    """渲染主页"""
    return render_template('index.html')

@app.route('/save_api_key', methods=['POST'])
def save_api_key():
    """保存API密钥到会话"""
    data = request.get_json()
    api_key = data.get('api_key', '')
    
    if api_key:
        session['api_key'] = api_key
        return jsonify({"status": "success", "message": "API密钥已保存"})
    else:
        return jsonify({"status": "error", "message": "API密钥不能为空"}), 400

@app.route('/upload_docx', methods=['POST'])
def upload_docx():
    """处理Word文档上传"""
    if 'docxFile' not in request.files:
        return jsonify({"status": "error", "message": "未找到文件"}), 400
    
    file = request.files['docxFile']
    
    if file.filename == '':
        return jsonify({"status": "error", "message": "未选择文件"}), 400
    
    if file and file.filename.endswith('.docx'):
        try:
            # 读取Word文档内容
            content = docx2txt.process(file)
            
            # 将内容保存到会话中
            session['uploaded_content'] = content
            
            return jsonify({
                "status": "success", 
                "message": "文档上传成功", 
                "content": content
            })
        except Exception as e:
            return jsonify({"status": "error", "message": f"处理文档时出错: {str(e)}"}), 500
    else:
        return jsonify({"status": "error", "message": "请上传.docx格式的文件"}), 400

@app.route('/download_docx', methods=['POST'])
def download_docx():
    """将生成的新闻稿下载为Word文档"""
    data = request.get_json()
    content = data.get('content', '')
    title = data.get('title', '新闻稿')
    
    if not content:
        return jsonify({"status": "error", "message": "内容不能为空"}), 400
    
    try:
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading(title, 0)
        
        # 添加内容（按段落分割）
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para)
        
        # 保存到内存中
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # 发送文件
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=f"{title}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({"status": "error", "message": f"创建Word文档时出错: {str(e)}"}), 500

@app.route('/generate_news', methods=['POST'])
def generate_news():
    """生成新闻稿"""
    # 获取API密钥
    api_key = session.get('api_key', '')
    if not api_key:
        return jsonify({"status": "error", "message": "请先设置API密钥"}), 400
    
    # 获取用户输入
    data = request.get_json()
    news_type = data.get('news_type', '')
    topic = data.get('topic', '')
    keywords = data.get('keywords', '')
    length = data.get('length', 'medium')
    tone = data.get('tone', 'neutral')
    reference_content = data.get('reference_content', '')
    
    # 验证输入
    if not topic:
        return jsonify({"status": "error", "message": "请输入新闻主题"}), 400
    
    # 构建提示词
    prompt = f"""请根据以下信息生成一篇新闻稿：
    - 新闻类型：{news_type}
    - 主题：{topic}
    - 关键词：{keywords}
    - 长度：{length}
    - 语调：{tone}
    """
    
    # 如果有参考内容，添加到提示词中
    if reference_content:
        prompt += f"""
    - 参考内容：
    {reference_content}
    
    请基于上述参考内容，生成一篇结构完整、内容丰富、语言流畅的新闻稿。包括标题、导语、主体内容和结尾。
    """
    else:
        prompt += """
    请生成一篇结构完整、内容丰富、语言流畅的新闻稿。包括标题、导语、主体内容和结尾。
    """
    
    try:
        # 调用DeepSeek API
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        
        payload = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.7,
            "max_tokens": 2000
        }
        
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload)
        response_data = response.json()
        
        # 处理API响应
        if response.status_code == 200 and "choices" in response_data:
            news_content = response_data["choices"][0]["message"]["content"]
            return jsonify({"status": "success", "news_content": news_content})
        else:
            error_message = response_data.get("error", {}).get("message", "未知错误")
            return jsonify({"status": "error", "message": f"API调用失败: {error_message}"}), 500
    
    except Exception as e:
        return jsonify({"status": "error", "message": f"发生错误: {str(e)}"}), 500

if __name__ == '__main__':
    app.run(debug=True)
